"""
Document Processor - extracts text content from uploaded files.

Supports: PDF, DOCX, TXT, MD, XLSX
Returns: structured Document objects with extracted text and metadata.
"""
import os
import csv
from pathlib import Path
from datetime import datetime
from typing import Optional
from pydantic import BaseModel, Field

# Document parsing libraries
import pypdf
from docx import Document as DocxDocument
from openpyxl import load_workbook
import pytesseract
from pdf2image import convert_from_path


# =============================================================================
# Document model - the structured output of ingestion
# =============================================================================
class IngestedDocument(BaseModel):
    """A document that has been processed and is ready for analysis."""
    
    filename: str = Field(
        description="Original filename"
    )
    file_type: str = Field(
        description="File extension (pdf, docx, txt, md, xlsx)"
    )
    file_size_bytes: int = Field(
        description="Size of the original file"
    )
    claimed_purpose: str = Field(
        default="",
        description="What the user said this document demonstrates"
    )
    extracted_text: str = Field(
        description="The full extracted text content"
    )
    page_count: Optional[int] = Field(
        default=None,
        description="Number of pages (PDF), paragraphs (DOCX), or sheets (XLSX)"
    )
    extraction_warnings: list[str] = Field(
        default_factory=list,
        description="Any issues encountered during extraction"
    )
    ingested_at: str = Field(
        description="ISO timestamp when this document was processed"
    )


# =============================================================================
# Format-specific extractors
# =============================================================================
def _ocr_pdf(file_path: str) -> tuple[str, int, list[str]]:
    """
    OCR a scanned PDF by converting pages to images and running Tesseract.
    
    Slower than text extraction (~5-10 sec per page) and quality varies with
    scan quality. Used only as a fallback when native text extraction fails.
    
    Returns (text, page_count, warnings).
    """
    warnings = []
    text_parts = []
    
    try:
        # Convert all PDF pages to PIL images at 200 DPI
        # 200 DPI is a good balance between OCR accuracy and processing speed
        images = convert_from_path(file_path, dpi=200)
        page_count = len(images)
        
        for page_num, image in enumerate(images, 1):
            try:
                ocr_text = pytesseract.image_to_string(image)
                if ocr_text and ocr_text.strip():
                    text_parts.append(f"[Page {page_num} (OCR)]\n{ocr_text.strip()}")
                else:
                    warnings.append(f"Page {page_num}: OCR found no text")
            except Exception as e:
                warnings.append(f"Page {page_num} OCR failed: {str(e)[:100]}")
        
        if not text_parts:
            warnings.append(
                "OCR found no extractable text. Document may be very low quality, "
                "blank, or contain only diagrams/images."
            )
    
    except Exception as e:
        raise RuntimeError(f"OCR processing failed: {e}")
    
    full_text = "\n\n".join(text_parts)
    return full_text, page_count, warnings

def _extract_pdf(file_path: str) -> tuple[str, int, list[str]]:
    """
    Extract text from a PDF file.
    
    Strategy: Try native text extraction first (fast, high quality).
    If that yields almost no text, the PDF is likely scanned -- fall back to OCR.
    
    Returns (text, page_count, warnings).
    """
    warnings = []
    text_parts = []
    pages_with_no_text = []
    
    try:
        reader = pypdf.PdfReader(file_path)
        page_count = len(reader.pages)
        
        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text.strip()}")
                else:
                    pages_with_no_text.append(page_num)
            except Exception as e:
                warnings.append(f"Page {page_num} extraction failed: {str(e)[:100]}")
    
    except Exception as e:
        raise RuntimeError(f"Could not open PDF: {e}")
    
    full_text = "\n\n".join(text_parts)
    
    # Decide whether to fall back to OCR
    # Threshold: if native extraction got <100 chars total, the PDF is likely scanned
    extraction_was_minimal = len(full_text.strip()) < 100
    
    if extraction_was_minimal:
        warnings.append(
            f"Native text extraction yielded only {len(full_text.strip())} chars; "
            f"falling back to OCR. This will take ~5-10 seconds per page."
        )
        try:
            ocr_text, ocr_page_count, ocr_warnings = _ocr_pdf(file_path)
            full_text = ocr_text
            page_count = ocr_page_count
            warnings.extend(ocr_warnings)
            warnings.append("Text extracted via OCR; some character errors may exist.")
        except Exception as e:
            warnings.append(f"OCR fallback also failed: {str(e)[:200]}")
    elif pages_with_no_text:
        # Some pages had no text but most were fine -- mention which ones
        warnings.append(
            f"Pages with no extractable text: {pages_with_no_text}. "
            f"These may be scanned/image pages within an otherwise text-based document."
        )
    
    return full_text, page_count, warnings


def _extract_docx(file_path: str) -> tuple[str, int, list[str]]:
    """Extract text from a Word document. Returns (text, paragraph_count, warnings)."""
    warnings = []
    
    try:
        doc = DocxDocument(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        for table_num, table in enumerate(doc.tables, 1):
            table_text_parts = [f"\n[Table {table_num}]"]
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text_parts.append(row_text)
            if len(table_text_parts) > 1:
                paragraphs.append("\n".join(table_text_parts))
        
        if not paragraphs:
            warnings.append("No text content found in document (might be image-based)")
    
    except Exception as e:
        raise RuntimeError(f"Could not open DOCX: {e}")
    
    full_text = "\n\n".join(paragraphs)
    paragraph_count = len(paragraphs)
    return full_text, paragraph_count, warnings


def _extract_text(file_path: str) -> tuple[str, int, list[str]]:
    """Extract text from a plain text or markdown file."""
    warnings = []
    
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            content = f.read()
        warnings.append("File was not UTF-8; read as latin-1")
    
    line_count = content.count("\n") + 1
    return content, line_count, warnings


def _extract_xlsx(file_path: str) -> tuple[str, int, list[str]]:
    """Extract text from an Excel file. Returns (text, sheet_count, warnings)."""
    warnings = []
    text_parts = []
    
    try:
        workbook = load_workbook(file_path, data_only=True, read_only=True)
        sheet_count = len(workbook.sheetnames)
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"\n[Sheet: {sheet_name}]")
            
            row_count = 0
            
            for row in sheet.iter_rows(values_only=True):
                if all(cell is None or str(cell).strip() == "" for cell in row):
                    continue
                
                row_text = " | ".join(
                    str(cell).strip() if cell is not None else ""
                    for cell in row
                )
                text_parts.append(row_text)
                row_count += 1
            
            if row_count == 0:
                warnings.append(f"Sheet '{sheet_name}': no data rows found")
        
        workbook.close()
    
    except Exception as e:
        raise RuntimeError(f"Could not open Excel file: {e}")
    
    if not text_parts:
        warnings.append("Excel file contains no data")
    
    full_text = "\n".join(text_parts)
    return full_text, sheet_count, warnings
def _extract_csv(file_path: str) -> tuple[str, int, list[str]]:
    """Extract text from a CSV file. Returns (text, row_count, warnings)."""
    warnings = []
    text_parts = []
    
    try:
        # Try UTF-8 first
        try:
            with open(file_path, "r", encoding="utf-8", newline="") as f:
                reader = csv.reader(f)
                rows = list(reader)
        except UnicodeDecodeError:
            # Fall back to latin-1 for files with weird encoding
            with open(file_path, "r", encoding="latin-1", newline="") as f:
                reader = csv.reader(f)
                rows = list(reader)
            warnings.append("File was not UTF-8; read as latin-1")
        
        for row in rows:
            # Skip rows that are entirely empty
            if not row or all(cell.strip() == "" for cell in row):
                continue
            
            row_text = " | ".join(cell.strip() for cell in row)
            text_parts.append(row_text)
        
        if not text_parts:
            warnings.append("CSV file contains no data rows")
    
    except Exception as e:
        raise RuntimeError(f"Could not parse CSV file: {e}")
    
    full_text = "\n".join(text_parts)
    row_count = len(text_parts)
    return full_text, row_count, warnings


# =============================================================================
# Main entry point
# =============================================================================
def ingest_document(
    file_path: str,
    claimed_purpose: str = "",
) -> IngestedDocument:
    """
    Process a document file and return an IngestedDocument with extracted text.
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")
    
    file_type = path.suffix.lower().lstrip(".")
    file_size = path.stat().st_size
    
    if file_type == "pdf":
        text, page_count, warnings = _extract_pdf(str(path))
    elif file_type == "docx":
        text, page_count, warnings = _extract_docx(str(path))
    elif file_type in ("txt", "md", "markdown"):
        text, page_count, warnings = _extract_text(str(path))
    elif file_type in ("xlsx", "xlsm"):
        text, page_count, warnings = _extract_xlsx(str(path))
    elif file_type == "csv":
        text, page_count, warnings = _extract_csv(str(path))
    else:
        raise ValueError(
            f"Unsupported file type: .{file_type}. "
            f"Supported: pdf, docx, txt, md, xlsx, csv"
        )
    
    return IngestedDocument(
        filename=path.name,
        file_type=file_type,
        file_size_bytes=file_size,
        claimed_purpose=claimed_purpose,
        extracted_text=text,
        page_count=page_count,
        extraction_warnings=warnings,
        ingested_at=datetime.now().isoformat(),
    )